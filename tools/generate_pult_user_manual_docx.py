# -*- coding: utf-8 -*-
"""
Генерация пользовательской справки по пульту агента (AIStudio + движок isida).
Запуск: python generate_pult_user_manual_docx.py

Зависимость: pip install python-docx
Результат сохраняется рядом со старым «Руководство.docx» под новым именем.
"""

from __future__ import annotations

import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


HELP_ROOT = Path(r"D:\ISIDA\help")
SOURCE_SUBDIR_NAME = "Справка юзера"
OLD_FILENAME = "Руководство.docx"
OUTPUT_FILENAME = "Руководство_пульт_AIStudio_2026.docx"


def find_help_dir() -> Path | None:
    if not HELP_ROOT.is_dir():
        return None
    for child in HELP_ROOT.iterdir():
        if child.is_dir() and child.name == SOURCE_SUBDIR_NAME:
            return child
    return None


def extract_plain_paragraphs_from_docx(path: Path) -> list[str]:
    """Читает текст из docx без python-docx (для сверки со старым стилем)."""
    paras: list[str] = []
    try:
        with zipfile.ZipFile(path, "r") as z:
            xml_data = z.read("word/document.xml").decode("utf-8")
    except (OSError, KeyError, zipfile.BadZipFile):
        return paras
    root = ET.fromstring(xml_data)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for p in root.findall(".//w:p", ns):
        texts = []
        for t in p.findall(".//w:t", ns):
            if t.text:
                texts.append(t.text)
            if t.tail:
                texts.append(t.tail)
        line = "".join(texts).strip()
        if line:
            paras.append(line)
    return paras


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    # Явно задаём шрифт для кириллицы в Word
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        p.paragraph_format.space_after = Pt(3)


def build_manual(doc: Document) -> None:
    set_document_defaults(doc)

    title = doc.add_paragraph()
    tr = title.add_run(
        "Работа с пультом взаимодействия с агентом (AIStudio / ISIDA)"
    )
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = "Times New Roman"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Краткое руководство для оператора AIStudio. Обновлено под текущую версию интерфейса и движка isida."
    )
    sr.font.size = Pt(12)
    sr.font.name = "Times New Roman"
    sr.italic = True
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_heading(doc, "Введение", 1)
    add_para(
        doc,
        "Пульт — это часть вкладки «Агент»: вы задаёте словесные стимулы и отмечаете воздействия на «жизненные параметры» модели. "
        "Настройка самих типов воздействий (названия, влияние на параметры, антагонисты) выполняется на других страницах раздела «Гомеостаз», "
        "а пульт служит для быстрой подачи уже подготовленных воздействий в ходе эксперимента.",
    )

    add_heading(doc, "1. Где находится пульт и зачем он нужен", 1)
    add_para(
        doc,
        "Пульт — это панель «Пульт взаимодействия с агентом» в разделе «Агент» главного окна AIStudio. "
        "Через него вы подаёте агенту словесные сообщения, отмечаете типы внешних воздействий на его «жизненные параметры», "
        "задаёте тон и настроение реплики, при необходимости — цвет сцены для зрительного канала. "
        "Ниже по той же вкладке показываются параметры агента и его действия; пульт работает вместе с остальной средой моделирования.",
    )

    add_heading(doc, "2. Пульсация — обязательное условие", 1)
    add_para(
        doc,
        "Все подачи с пульта привязаны к работе таймера пульсации («пульсы»). Если пульсация выключена, "
        "кнопка «Отправить» не применит воздействия — программа сообщит об этом. Сценарии оператора и автоматическая подача стимулов тоже требуют включённой пульсации.",
    )

    add_heading(doc, "3. Поле «Окно сообщения агенту» и распознавание текста", 1)
    add_para(
        doc,
        "Вы вводите фразу так, как хотите её «сказать» агенту. Справа от поля показывается «Распознанный текст»: "
        "слова, которые уже есть в словарном дереве агента, отображаются как есть; слова, которых в дереве ещё нет, "
        "заменяются на группы символов «xxxxx». Это не ошибка ввода — так отмечаются пока неизвестные слова.",
    )
    add_para(
        doc,
        "После нажатия «Отправить» текстовое поле очищается; сам стимул уходит в модель вместе с выбранными галочками воздействий, тоном, настроением и цветом.",
    )

    add_heading(doc, "4. Режим наблюдения", 1)
    add_para(
        doc,
        "Если включён флажок «Режим наблюдения», воздействия на жизненные параметры гомеостаза не меняют числовые значения — "
        "как будто агент «видит» действие со стороны, без реального толчка по параметрам. При этом события восприятия и запись в память по фразе могут продолжаться — "
        "режим нужен, чтобы отработать сценарий «наблюдателя», не искажая показатели здоровья агента.",
    )

    add_heading(doc, "5. Авторитарная запись", 1)
    add_para(
        doc,
        "Если включена «Авторитарная запись», новые слова из вашей фразы попадают в дерево слов сразу, минуя промежуточную «песочницу». "
        "Если выключена — используется обычный путь записи (по правилам проекта). Используйте осознанно: быстрая запись ускоряет обучение словарю, но снижает контроль.",
    )

    add_heading(doc, "6. Тон, настроение и цвет", 1)
    add_para(
        doc,
        "«Тон» и «Настроение» задают эмоциональную окраску реплики при распознавании текста — они участвуют в образе фразового стимула.",
    )
    add_para(
        doc,
        "«Цвет» — это код зрительного канала фона (по умолчанию белый). Ненулевой цвет задаёт «цветовую сцену» для пускового образа; "
        "для составных условных связок может использоваться сочетание фразы и цвета. После успешной отправки цвет обычно сбрасывается на белый.",
    )

    add_heading(doc, "7. Списки воздействий (галочки)", 1)
    add_para(
        doc,
        "В двух столбцах перечислены воздействия из файла проекта «InfluenceActions» — каждое связано с изменением тех или иных жизненных параметров агента. "
        "Точные названия и смысл берутся из настроек вашего проекта; при наведении на галочку показывается подсказка с описанием.",
    )
    add_para(
        doc,
        "Можно отметить несколько воздействий одновременно; некоторые сочетания могут быть связаны как «антагонисты» (взаимоисключающие воздействия) — логика подсказок также задаётся в проекте.",
    )

    add_heading(doc, "8. Кнопка «Отправить» и клавиша Enter", 1)
    add_para(
        doc,
        "«Отправить» применяет текущий набор: отмеченные воздействия, текст (если есть), тон, настроение и цвет. "
        "Если не выбрано ни одного воздействия, поле текста пусто и цвет остаётся белым, отправить нечего — программа сообщит об этом.",
    )
    add_para(
        doc,
        "В области пульта клавиша Enter дублирует «Отправить» (если фокус не в выпадающем списке и не в многострочном поле с переносами).",
    )

    add_heading(doc, "9. Цепочки рефлексов и автоматизмов", 1)
    add_para(
        doc,
        "Когда выполняется цепочка рефлексов, появляется цветная плашка «Цепочка активна» и переключатель «Результат звена: Успех / Неудача». "
        "Вы сообщаете системе, как завершилось текущее действие в цепи — это влияет на развитие связки.",
    )
    add_para(
        doc,
        "Для цепочки автоматизмов вместо оценочной плашки показывается текст вида «Выполняется звено цепочки №…» — оценка шага может опираться на эффект только что поданного стимула.",
    )

    add_heading(doc, "10. Плашка «ожидание оценки оператора»", 1)
    add_para(
        doc,
        "Над блоком действий агента может отображаться полоска с таймером ожидания. В этом режиме система ждёт вашего решения или завершения процедуры оценки. "
        "Клик по плашке или кнопка «✕» отменяют ожидание (как указано во всплывающей подсказке). Это связано с циклом «оператор оценил результат», а не с самим пультом.",
    )

    add_heading(doc, "11. Сценарии оператора и пульт", 1)
    add_para(
        doc,
        "При запуске сценария из раздела «Исследования» программа сама подаёт стимулы в нужные моменты пульса через тот же механизм, что и ручная отправка. "
        "Для этого нужно открыть вкладку «Агент», чтобы объект пульта был создан — иначе сценарий не сможет подать воздействия.",
    )
    add_para(
        doc,
        "Текстовые шаги сценария распознаются как реплики с принудительным режимом записи слов (аналог «авторитарной» подачи фразы), независимо от того, включён ли у вас на пульте флажок «Авторитарная запись».",
    )
    add_para(
        doc,
        "В заголовке сценария задаются режимы «Режим наблюдения» и «Авторитарная запись» на время прогона; после завершения сценария восстанавливаются ваши прежние установки на пульте.",
    )
    add_para(
        doc,
        "Шаги сценария бывают двух видов: с реальной подачей стимула (фраза, воздействия, цвет) и специальные шаги «ожидание клика» — они сбрасывают период ожидания оценки в нужный пульс без нового стимула.",
    )

    add_heading(doc, "12. Если агент «мёртв»", 1)
    add_para(
        doc,
        "При наступлении условий смерти агента пульт блокируется: воздействия не подаются, пока вы не восстановите состояние через настройки проекта или новый прогон по правилам вашей методики.",
    )

    doc.add_paragraph()
    add_para(
        doc,
        "Примечание: имена кнопок воздействий, состав параметров гомеостаза и словарь берутся из открытого проекта; это руководство описывает только поведение интерфейса пульта.",
        bold=False,
    )


def main() -> None:
    help_dir = find_help_dir()
    if help_dir is None:
        raise SystemExit(f"Не найден каталог {HELP_ROOT / SOURCE_SUBDIR_NAME}")

    old_path = help_dir / OLD_FILENAME
    legacy_note = ""
    if old_path.is_file():
        old_paras = extract_plain_paragraphs_from_docx(old_path)
        legacy_note = f" (фрагментов текста в старом файле: {len(old_paras)})"
        # Для отладки можно распечатать первые строки:
        # print(old_paras[:30])

    doc = Document()
    build_manual(doc)

    out_path = help_dir / OUTPUT_FILENAME
    doc.save(out_path)
    print(f"Записано: {out_path}{legacy_note}")


if __name__ == "__main__":
    main()
