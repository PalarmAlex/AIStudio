# -*- coding: utf-8 -*-
"""
Сборка справки Word из D:\\ISIDA\\help\\Справка юзера\\index.php (содержимое heredoc $str).

Зависимости:
  pip install python-docx beautifulsoup4

Запуск:
  python generate_help_docx_from_index_php.py
"""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

INDEX_PHP = Path(r"D:\ISIDA\help\Справка юзера\index.php")
OUTPUT_DOCX = Path(r"D:\ISIDA\help\Справка юзера\Руководство_ISIDA_из_index_2026.docx")


def ensure_bs4() -> None:
    try:
        import bs4  # noqa: F401
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "beautifulsoup4", "-q"])


def extract_heredoc_source(php_path: Path) -> str:
    raw = php_path.read_text(encoding="utf-8")
    m = re.search(r"\$str\s*=\s*<<<EOD\s*\r?\n(.*)\r?\nEOD\s*;", raw, re.DOTALL)
    if not m:
        raise ValueError("Не найден блок $str = <<<EOD ... EOD в index.php")
    return m.group(1).strip()


def set_run_font(run, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    try:
        rpr = run._element.rPr
        if rpr is not None and rpr.rFonts is not None:
            rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except (AttributeError, TypeError):
        pass


def add_paragraph_runs(doc: Document, element: Tag, *, italic_block: bool = False) -> None:
    """Текст из inline-тегов (strong, em, a, code) в один абзац."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def walk(node, bold: bool, italic: bool) -> None:
        if isinstance(node, NavigableString):
            t = str(node)
            if not t:
                return
            run = p.add_run(t)
            set_run_font(run, bold=bold, italic=italic or italic_block)
            return
        if not isinstance(node, Tag):
            return
        name = node.name.lower()
        if name == "br":
            return
        nb, ni = bold, italic
        if name in ("strong", "b"):
            nb = True
        elif name in ("em", "i"):
            ni = True
        elif name == "a":
            text = node.get_text(strip=True)
            href = node.get("href") or ""
            if href:
                text = f"{text} ({href})"
            run = p.add_run(text)
            set_run_font(run, bold=nb, italic=ni or italic_block)
            return
        elif name == "code":
            for child in node.children:
                walk(child, bold, True)
            return
        for child in node.children:
            walk(child, nb, ni)

    for child in element.children:
        walk(child, False, False)

    if not p.text.strip():
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)


def flush_children(doc: Document, node: Tag, *, skip_tags: frozenset[str] | None = None) -> None:
    skip = skip_tags or frozenset()
    for child in node.children:
        if isinstance(child, NavigableString):
            t = str(child).strip()
            if t:
                p = doc.add_paragraph(t)
                p.paragraph_format.space_after = Pt(6)
                for r in p.runs:
                    set_run_font(r)
            continue
        if not isinstance(child, Tag):
            continue
        process_block(doc, child, skip_tags=skip)


def process_block(doc: Document, el: Tag, *, skip_tags: frozenset[str] | None = None) -> None:
    skip = skip_tags or frozenset()
    name = el.name.lower()

    if name in skip:
        return

    if name == "div":
        cls_list = el.get("class") or []
        cls = " ".join(cls_list).lower()

        if "note" in cls:
            handled = False
            for ch in el.children:
                if isinstance(ch, NavigableString):
                    continue
                if not isinstance(ch, Tag):
                    continue
                if ch.name == "strong":
                    p = doc.add_paragraph()
                    r = p.add_run(ch.get_text(strip=True))
                    set_run_font(r, bold=True)
                    p.paragraph_format.space_after = Pt(6)
                    handled = True
                elif ch.name == "p":
                    add_paragraph_runs(doc, ch)
                    handled = True
            if not handled:
                add_paragraph_runs(doc, el)
            return

        if "warning" in cls:
            handled = False
            for ch in el.children:
                if isinstance(ch, NavigableString):
                    continue
                if not isinstance(ch, Tag):
                    continue
                if ch.name == "strong":
                    p = doc.add_paragraph()
                    r = p.add_run(ch.get_text(strip=True))
                    set_run_font(r, bold=True)
                    p.paragraph_format.space_after = Pt(6)
                    handled = True
                elif ch.name == "p":
                    add_paragraph_runs(doc, ch)
                    handled = True
            if not handled:
                add_paragraph_runs(doc, el)
            return

        if "principle-box" in cls:
            for ch in el.children:
                if not isinstance(ch, Tag):
                    continue
                ch_cls = ch.get("class") or []
                if ch.name == "div" and "principle-title" in ch_cls:
                    p = doc.add_paragraph()
                    r = p.add_run(ch.get_text(" ", strip=True))
                    set_run_font(r, bold=True)
                    p.paragraph_format.space_after = Pt(6)
                elif ch.name == "p":
                    add_paragraph_runs(doc, ch)
            return

        if "section" in cls or not cls_list:
            flush_children(doc, el)
            return

        flush_children(doc, el)
        return

    if name == "h2":
        t = el.get_text(" ", strip=True)
        doc.add_heading(t, level=1)
        return

    if name == "h3":
        t = el.get_text(" ", strip=True)
        doc.add_heading(t, level=2)
        return

    if name == "h4":
        t = el.get_text(" ", strip=True)
        doc.add_heading(t, level=3)
        return

    if name == "p":
        # параграф может содержать img — отдельно
        imgs = el.find_all("img")
        if imgs:
            for img in imgs:
                alt = img.get("alt") or "(иллюстрация)"
                cap = doc.add_paragraph(f"[Иллюстрация: {alt}]")
                cap.paragraph_format.space_after = Pt(4)
                for r in cap.runs:
                    set_run_font(r, italic=True)
        text_only = "".join(
            str(c) for c in el.contents if not (isinstance(c, Tag) and c.name == "img")
        )
        soup_mini = BeautifulSoup(f"<div>{text_only}</div>", "html.parser")
        div = soup_mini.div
        if div and div.get_text(strip=True):
            add_paragraph_runs(doc, div)
        return

    if name == "ul":
        for li in el.find_all("li", recursive=False):
            text = li.get_text(" ", strip=True)
            # вложенные ul обрабатываем упрощённо
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            set_run_font(run)
            p.paragraph_format.space_after = Pt(3)
        return

    if name == "ol":
        for li in el.find_all("li", recursive=False):
            text = li.get_text(" ", strip=True)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text)
            set_run_font(run)
            p.paragraph_format.space_after = Pt(3)
        return

    if name == "table":
        rows = el.find_all("tr")
        if not rows:
            return
        cols_max = max(len(r.find_all(["th", "td"])) for r in rows)
        ncols = max(cols_max, 1)
        tbl = doc.add_table(rows=len(rows), cols=ncols)
        tbl.style = "Table Grid"
        for ri, tr in enumerate(rows):
            cells = tr.find_all(["th", "td"])
            for ci in range(ncols):
                txt = cells[ci].get_text(" ", strip=True) if ci < len(cells) else ""
                tbl.rows[ri].cells[ci].text = txt
        doc.add_paragraph()
        return

    if name == "img":
        alt = el.get("alt") or "(иллюстрация)"
        cap = doc.add_paragraph(f"[Иллюстрация: {alt}]")
        for r in cap.runs:
            set_run_font(r, italic=True)
        return

    if name in ("style", "script"):
        return

    flush_children(doc, el)


def strip_toc_and_styles(html: str) -> str:
    """Убираем оглавление и блок style — они уже не нужны в Word."""
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all("style"):
        tag.decompose()
    # первый div.section с toc-list — удалить
    for div in soup.find_all("div", class_=True):
        cls = div.get("class") or []
        if "section" in cls:
            ul = div.find("ul", class_=lambda x: x and "toc-list" in x)
            if ul:
                div.decompose()
                break
    return str(soup)


def build_document(html_fragment: str) -> Document:
    html_fragment = strip_toc_and_styles(html_fragment)
    wrapped = f"<html><body>{html_fragment}</body></html>"
    soup = BeautifulSoup(wrapped, "html.parser")
    body = soup.body
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    title = doc.add_paragraph()
    tr = title.add_run(
        "Проектирование агентов с индивидуальной адаптивной архитектурой в системе ISIDA"
    )
    set_run_font(tr, bold=True)
    tr.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Текст подготовлен из локальной справки (index.php). Дата сборки: май 2026."
    )
    set_run_font(sr, italic=True)
    sr.font.size = Pt(11)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for child in body.children:
        if isinstance(child, NavigableString):
            continue
        if isinstance(child, Tag):
            process_block(doc, child)

    return doc


def main() -> None:
    ensure_bs4()
    if not INDEX_PHP.is_file():
        raise SystemExit(f"Не найден файл: {INDEX_PHP}")

    html = extract_heredoc_source(INDEX_PHP)
    doc = build_document(html)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Записано: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
